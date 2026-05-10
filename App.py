import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from groq import Groq
import json
import io

# ==========================================
# 1. WEB APP CONFIG & DARK AESTHETIC CSS
# ==========================================
st.set_page_config(page_title="OceanTailor AI v2.5", layout="wide", page_icon="🌊")

st.markdown("""
    <style>
    .stApp { background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%); color: #f1f5f9; }
    @keyframes fadeInUp { from { opacity: 0; transform: translateY(20px); } to { opacity: 1; transform: translateY(0); } }
    .main-container { animation: fadeInUp 0.8s ease-out; }
    h1, h2, h3 { color: #22d3ee !important; font-family: 'Inter', sans-serif; font-weight: 700 !important; }
    .profile-card { 
        padding: 25px; background: rgba(255, 255, 255, 0.05); backdrop-filter: blur(10px);
        border-radius: 20px; border: 1px solid rgba(255, 255, 255, 0.1);
        border-left: 6px solid #22d3ee; margin-bottom: 25px; 
    }
    .metric-card {
        background: rgba(255, 255, 255, 0.03); padding: 20px; border-radius: 15px;
        border: 1px solid rgba(34, 211, 238, 0.2); text-align: center; transition: all 0.3s ease;
    }
    .metric-card:hover { transform: translateY(-5px); border: 1px solid #22d3ee; box-shadow: 0px 10px 20px rgba(34, 211, 238, 0.1); }
    .stButton>button { 
        background: linear-gradient(90deg, #0891b2 0%, #22d3ee 100%); color: #0f172a !important; 
        border-radius: 12px; border: none; font-weight: bold; width: 100%; transition: all 0.3s ease;
    }
    .stButton>button:hover { transform: scale(1.02); box-shadow: 0px 0px 15px rgba(34, 211, 238, 0.6); }
    .stTextInput>div>div>input, .stTextArea>div>div>textarea { 
        background-color: rgba(15, 23, 42, 0.6) !important; color: #f1f5f9 !important;
        border: 1px solid rgba(34, 211, 238, 0.3) !important; border-radius: 12px !important; 
    }
    [data-testid="stSidebar"] { background-color: #0f172a !important; }
    </style>
    <div class="main-container"></div>
    """, unsafe_allow_html=True)

# ==========================================
# 2. LOGIC FUNCTIONS
# ==========================================
def extract_text_from_pdf(file):
    reader = PdfReader(file)
    return "".join([page.extract_text() for page in reader.pages])

def create_docx(text):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        if line.startswith('# '):
            heading = doc.add_heading(line.replace('# ', ''), level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = heading.runs[0]
            run.font.color.rgb = RGBColor(15, 23, 42)
            run.font.size = Pt(14)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line.replace('## ', ''))
            run.bold = True
            run.font.size = Pt(12)
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line.replace('* ', '').replace('- ', ''), style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
        else:
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def analyze_resume(base_text, jd, api_key):
    client = Groq(api_key=api_key)
    prompt = f"Analyze Resume vs JD. Return ONLY JSON: {{'match_score': int, 'missing_keywords': [], 'improvement_tips': []}}. JD: {jd} Resume: {base_text}"
    completion = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[{"role": "system", "content": "You are an ATS expert. Return JSON only."},
                  {"role": "user", "content": prompt}],
        response_format={"type": "json_object"}
    )
    return json.loads(completion.choices[0].message.content)

def tailor_resume(base_text, company, jd, api_key):
    client = Groq(api_key=api_key)
    prompt = f"""
    You are an expert ATS resume writer.
    Company: {company}
    JD: {jd}
    Base Resume: {base_text}
    TASK: Rewrite the resume. Use '# ' for Main Sections, '## ' for Job Titles, and '* ' for bullets.
    Maintain EXACT same structure. Use Google XYZ formula. Return ONLY the final text.
    """
    completion = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[{"role": "system", "content": "You are a world-class resume writer."},
                  {"role": "user", "content": prompt}]
    )
    return completion.choices[0].message.content

# ==========================================
# 3. PROFILE MANAGEMENT
# ==========================================
if 'profiles' not in st.session_state:
    st.session_state.profiles = {}

# ==========================================
# 4. UI LAYOUT
# ==========================================
st.title("🌊 OceanTailor AI v2.5")
st.markdown("#### *Professional Automated Edition*")

# --- PERMANENT API KEY IMPLEMENTATION ---
if "GROQ_API_KEY" in st.secrets:
    api_key = st.secrets["GROQ_API_KEY"]
else:
    st.error("❌ API Key missing! Please add 'GROQ_API_KEY' to your Streamlit Cloud Secrets.")
    st.stop()

with st.sidebar:
    st.header("👤 User Profiles")
    new_p = st.text_input("Profile Name")
    if st.button("➕ Create Profile"):
        if new_p:
            st.session_state.profiles[new_p] = ""
            st.success(f"Profile '{new_p}' Created!")
    st.divider()
    profile_list = list(st.session_state.profiles.keys())
    active_profile = st.selectbox("Select Active Profile", profile_list) if profile_list else None

if active_profile:
    st.markdown(f"<div class='profile-card'><b>Active Profile:</b> {active_profile}</div>", unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("📄 Base Resume")
        up_file = st.file_uploader("Upload PDF", type="pdf")
        if up_file:
            st.session_state.profiles[active_profile] = extract_text_from_pdf(up_file)
            st.success("Resume Loaded!")

    with col2:
        st.subheader("🎯 Job Target")
        comp = st.text_input("Company Name")
        jd = st.text_area("Job Description", height=200)

    st.divider()
    
    btn_col1, btn_col2 = st.columns(2)
    with btn_col1:
        if st.button("🔍 Analyze Match Score"):
            if not st.session_state.profiles[active_profile] or not jd:
                st.error("Missing Resume or JD!")
            else:
                with st.spinner("Analyzing..."):
                    try:
                        st.session_state.analysis = analyze_resume(st.session_state.profiles[active_profile], jd, api_key)
                    except Exception as e:
                        st.error(f"Error: {e}")

    with btn_col2:
        if st.button("🚀 Tailor My Resume"):
            if not st.session_state.profiles[active_profile] or not jd:
                st.error("Missing Resume or JD!")
            else:
                with st.spinner("Crafting..."):
                    try:
                        st.session_state.final_text = tailor_resume(st.session_state.profiles[active_profile], comp, jd, api_key)
                    except Exception as e:
                        st.error(f"Error: {e}")

    if 'analysis' in st.session_state:
        st.subheader("📊 ATS Analysis")
        res = st.session_state.analysis
        m_col1, m_col2, m_col3 = st.columns(3)
        with m_col1:
            st.markdown(f"<div class='metric-card'><h3>Match Score</h3><h2 style='color:#22d3ee'>{res['match_score']}%</h2></div>", unsafe_allow_html=True)
        with m_col2:
            st.markdown(f"<div class='metric-card'><h3>Missing Keywords</h3><p>{', '.join(res['missing_keywords'])}</p></div>", unsafe_allow_html=True)
        with m_col3:
            st.markdown(f"<div class='metric-card'><h3>Quick Tip</h3><p>{res['improvement_tips'][0]}</p></div>", unsafe_allow_html=True)

    if 'final_text' in st.session_state:
        st.divider()
        st.subheader("✨ Your Tailored Resume Preview")
        st.markdown("---")
        st.markdown(st.session_state.final_text)
        st.markdown("---")
        docx_b = create_docx(st.session_state.final_text)
        st.download_button("📥 Download Professional Word (.docx)", data=docx_b, file_name=f"Tailored_{comp}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
else:
    st.info("👈 Please create or select a profile from the sidebar to begin.")
